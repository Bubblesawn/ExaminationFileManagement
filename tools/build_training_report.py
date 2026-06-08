from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = Path.home() / "Downloads"
TEMPLATE = next(DOWNLOADS.glob("10.*实习报告.docx"))
OUTPUT = ROOT / "docs" / "省考试院自学考试考籍管理系统实训报告.docx"
ASSET_DIR = ROOT / "target" / "report-assets"

SCREENSHOT_MAP = {
    "系统登录页面": "12-login.png",
    "系统工作台页面": "02-dashboard.png",
    "登录认证流程": "12-login.png",
    "工作台统计卡片": "02-dashboard.png",
    "考生与考籍档案列表": "10-records.png",
    "材料审核页面": "03-material-audit.png",
    "免考申请详情": "04-exemption.png",
    "课程顶替规则与申请": "05-course-replace.png",
    "考籍转入转出审核": "06-transfer.png",
    "毕业资格校验": "07-graduation.png",
    "智能辅助页面": "08-ai-assistant.png",
    "日志管理页面": "09-logs.png",
    "系统测试用例执行结果": "03-material-audit.png",
}


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_style(paragraph, size: float = 11, bold: bool = False) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 10.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "9E9E9E")


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True)
    r.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc: Document, text: str, size: float = 11, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    set_paragraph_style(p, size=size, bold=bold)


def add_bullets(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"第{index}项：{item}")
        set_run_font(r, size=10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], "F2F2F2")
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=10)
    doc.add_paragraph()


def add_screenshot_placeholder(doc: Document, title: str, note: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    cell.height = Cm(5.8)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    image_name = SCREENSHOT_MAP.get(title)
    image_path = ASSET_DIR / image_name if image_name else None
    if image_path and image_path.exists():
        r = p.add_run()
        r.add_picture(str(image_path), width=Cm(12.6))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Pt(0)
        cr = cap.add_run(f"图：{title}")
        set_run_font(cr, size=9.5, bold=True)
    else:
        p.paragraph_format.space_before = Pt(72)
        r = p.add_run(f"【截图占位符：{title}】\n{note}\n请在此处替换为实际系统截图")
        set_run_font(r, size=12, bold=True)
        r.font.color.rgb = RGBColor(90, 90, 90)
    doc.add_paragraph()


def font(size: int, bold: bool = False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str, size: int = 28, bold: bool = False) -> None:
    f = font(size, bold)
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=f)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + (len(lines) - 1) * 8
    y = box[1] + ((box[3] - box[1]) - total_h) / 2
    for i, line in enumerate(lines):
        x = box[0] + ((box[2] - box[0]) - widths[i]) / 2
        draw.text((x, y), line, font=f, fill=fill)
        y += heights[i] + 8


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#51606F", width: int = 4) -> None:
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - direction * 16, y2 - 9), (x2 - direction * 16, y2 + 9)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - direction * 16), (x2 + 9, y2 - direction * 16)]
    draw.polygon(points, fill=color)


def make_diagrams() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    w, h = 1600, 900

    def base(title: str):
        img = Image.new("RGB", (w, h), "#FFFFFF")
        d = ImageDraw.Draw(img)
        d.rectangle((0, 0, w, 86), fill="#1F4E79")
        d.text((50, 23), title, font=font(34, True), fill="#FFFFFF")
        return img, d

    img, d = base("系统总体架构图")
    layers = [
        ((90, 170, 360, 300), "前端展示层\nVue3 + Element Plus", "#EAF3FF"),
        ((470, 170, 740, 300), "后端业务层\nSpring Boot", "#EAF8F0"),
        ((850, 170, 1120, 300), "算法服务层\nFastAPI", "#FFF4E5"),
        ((1230, 170, 1500, 300), "数据持久层\nMySQL + 文件存储", "#F4F0FF"),
        ((470, 460, 740, 610), "认证权限\nToken / Role / Menu", "#F7FBFF"),
        ((850, 460, 1120, 610), "智能辅助\n材料预处理 / 问答", "#FFF9EC"),
    ]
    for box, text, fill in layers:
        d.rounded_rectangle(box, radius=18, fill=fill, outline="#9AA7B4", width=3)
        draw_center(d, box, text, "#16324F", 28, True)
    arrow(d, (360, 235), (470, 235))
    arrow(d, (740, 235), (850, 235))
    arrow(d, (1120, 235), (1230, 235))
    arrow(d, (605, 300), (605, 460))
    arrow(d, (985, 300), (985, 460))
    draw_center(d, (270, 700, 1330, 790), "浏览器页面通过 /api 访问后端，后端统一校验 Token、处理业务规则，并以内网密钥调用算法服务。", "#333333", 28)
    img.save(ASSET_DIR / "diagram-architecture.png")

    img, d = base("业务办理流程图")
    steps = [
        ((80, 210, 290, 330), "登录系统"),
        ((360, 210, 570, 330), "选择业务\n免考/顶替/转考/毕业"),
        ((640, 210, 850, 330), "填写申请\n上传材料"),
        ((920, 210, 1130, 330), "智能核验\n材料预审"),
        ((1200, 210, 1510, 330), "人工审核\n通过/驳回"),
    ]
    for box, text in steps:
        d.rounded_rectangle(box, radius=18, fill="#EEF6FF", outline="#5B8CC0", width=3)
        draw_center(d, box, text, "#16324F", 26, True)
    for i in range(len(steps) - 1):
        arrow(d, (steps[i][0][2], 270), (steps[i + 1][0][0], 270))
    lower = [
        ((360, 520, 570, 650), "写入\nbusiness_application"),
        ((640, 520, 850, 650), "写入\nrecord_material"),
        ((920, 520, 1130, 650), "写入\naudit_record"),
        ((1200, 520, 1510, 650), "联动档案状态\n转出/毕业"),
    ]
    for box, text in lower:
        d.rounded_rectangle(box, radius=18, fill="#F6F8FA", outline="#A7B0BA", width=3)
        draw_center(d, box, text, "#333333", 24, True)
    arrow(d, (465, 330), (465, 520))
    arrow(d, (745, 330), (745, 520))
    arrow(d, (1025, 330), (1025, 520))
    arrow(d, (1355, 330), (1355, 520))
    img.save(ASSET_DIR / "diagram-business-flow.png")

    img, d = base("申请状态流转图")
    boxes = {
        "SUBMITTED\n已提交": (120, 270, 390, 400),
        "APPROVED\n审核通过": (520, 120, 790, 250),
        "REJECTED\n审核驳回": (520, 420, 790, 550),
        "WITHDRAWN\n已撤回": (930, 270, 1200, 400),
        "档案状态联动\nTRANSFERRED_OUT / GRADUATED": (930, 120, 1480, 250),
    }
    for text, box in boxes.items():
        fill = "#EAF8F0" if "APPROVED" in text else "#FFF4E5" if "REJECTED" in text else "#F4F0FF" if "WITHDRAWN" in text else "#EAF3FF"
        d.rounded_rectangle(box, radius=18, fill=fill, outline="#8D99A6", width=3)
        draw_center(d, box, text, "#16324F", 26, True)
    arrow(d, (390, 335), (520, 185))
    arrow(d, (390, 335), (520, 485))
    arrow(d, (390, 335), (930, 335))
    arrow(d, (790, 185), (930, 185))
    d.text((230, 205), "审核通过", font=font(22, True), fill="#3D6E42")
    d.text((255, 445), "审核驳回", font=font(22, True), fill="#8A5A00")
    d.text((610, 300), "撤回申请", font=font(22, True), fill="#5A4A7A")
    d.text((1040, 470), "终态不允许重复修改或审核", font=font(28, True), fill="#9B1C1C")
    img.save(ASSET_DIR / "diagram-status-flow.png")

    img, d = base("核心数据关系图")
    ers = [
        ((80, 165, 330, 275), "candidate\n考生"),
        ((430, 165, 720, 275), "student_record\n考籍档案"),
        ((820, 165, 1110, 275), "business_application\n业务申请"),
        ((1210, 165, 1510, 275), "audit_record\n审核记录"),
        ((430, 455, 720, 565), "record_material\n档案材料"),
        ((820, 455, 1110, 565), "application_extension_field\n扩展字段"),
        ((1210, 455, 1510, 565), "record_status_log\n状态日志"),
    ]
    for box, text in ers:
        d.rounded_rectangle(box, radius=14, fill="#F8FAFC", outline="#607D9B", width=3)
        draw_center(d, box, text, "#16324F", 23, True)
    arrow(d, (330, 220), (430, 220))
    arrow(d, (720, 220), (820, 220))
    arrow(d, (1110, 220), (1210, 220))
    arrow(d, (575, 275), (575, 455))
    arrow(d, (965, 275), (965, 455))
    arrow(d, (965, 275), (1210, 510))
    d.text((355, 185), "1:N", font=font(20, True), fill="#555555")
    d.text((745, 185), "1:N", font=font(20, True), fill="#555555")
    d.text((1135, 185), "1:N", font=font(20, True), fill="#555555")
    img.save(ASSET_DIR / "diagram-er.png")


def add_figure(doc: Document, image_name: str, caption: str, width_cm: float = 13.0) -> None:
    image_path = ASSET_DIR / image_name
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    r = cap.add_run(caption)
    set_run_font(r, size=9.5, bold=True)


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Pt(0)
    for line in text.splitlines():
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(9.5)
        r.add_break()
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("        计算机技术综合项目实践                 ")
    set_run_font(r, size=18, bold=True)
    for _ in range(3):
        doc.add_paragraph()
    lines = [
        ("设计题目", "省考试院自学考试考籍管理系统"),
        ("学院名称", "计算机与网络安全学院（示范性软件学院）"),
        ("专业名称", "计算机科学与技术"),
        ("学生姓名", "第 9 组"),
        ("学生学号", "（请在此处填写成员学号）"),
        ("任课教师", "张益恒、陈光柱"),
        ("设计（论文）成绩", ""),
    ]
    for label, value in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(2.3)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(f"{label} \t{value}")
        set_run_font(r, size=14)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("教务处 制")
    set_run_font(r, size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年 5月31日")
    set_run_font(r, size=12)
    add_page_break(doc)


def add_responsibility(doc: Document) -> None:
    add_heading(doc, "一、职责描述", 1)
    add_para(doc, "本项目以省考试院自学考试考籍管理业务为对象，围绕考生信息维护、考籍档案管理、材料审核、免考申请、课程顶替、考籍转入转出、毕业申请、日志审计和智能辅助核验等场景，完成了一个前后端分离、三端联调、具备基础智能处理能力的业务管理系统。团队在实训过程中按照需求分析、数据库建模、接口设计、前端实现、算法服务封装、联调测试和文档整理的顺序推进，形成了可运行、可演示、可继续扩展的课程项目。")
    add_table(
        doc,
        ["小组成员", "负责板块"],
        [
            ["成员一", "负责登录认证、权限控制、工作台统计、系统用户与角色菜单管理等后台基础能力，实现 Token 校验、权限标识控制和前端路由守卫。"],
            ["成员二", "负责考生信息、考籍档案、档案状态变更、归档和变更记录等核心业务模块，实现考籍数据的增删改查和状态追溯。"],
            ["成员三", "负责材料管理、材料上传预览、材料审核和业务材料绑定，实现材料与免考、顶替、转考、毕业申请的关联。"],
            ["成员四", "负责免考申请、课程顶替、转入转出、毕业申请等流程业务，实现提交、修改、撤回、审核、驳回与流程记录查询。"],
            ["成员五", "负责算法服务、智能问答、图像分类、材料预处理和申请材料智能核验，实现 Spring Boot 后端与 FastAPI 服务的安全联调。"],
            ["全体成员", "共同完成数据库初始化脚本、接口说明文档、系统测试用例、部署说明和最终演示材料整理。"],
        ],
    )
    add_para(doc, "本人在实训中重点参与系统整体方案梳理、后端接口实现、前端页面联调、算法服务对接、数据库脚本维护和报告文档编写。具体工作包括：根据自学考试考籍管理业务抽象数据模型；设计通用业务申请表与审核记录表；实现业务状态流转与档案状态联动；编写材料审核页面和智能核验接口说明；整理部署、测试和验收文档；在联调阶段记录问题并逐项修复。")
    add_bullets(
        doc,
        [
            "需求分析阶段：梳理省考试院自学考试考籍管理的业务角色、业务对象、数据关系和主要办理流程，明确系统需要支持的高频业务。",
            "设计阶段：完成前后端分离架构设计、数据库实体关系设计、接口路径规范、统一响应格式和权限标识设计。",
            "实现阶段：参与 Spring Boot 控制层、服务层、Mapper 层开发，配合 Vue3 页面完成接口联调，并对关键类补充 Doxygen 风格注释。",
            "测试阶段：根据系统功能清单执行登录、权限、档案、材料、流程、智能辅助和异常场景测试，保证主流程可连续走通。",
            "文档阶段：整理接口说明、部署说明、数据库设计和实训报告，使项目交付物能够支撑演示、答辩和后续维护。",
        ],
    )
    add_page_break(doc)


def add_design_description(doc: Document) -> None:
    add_heading(doc, "二、设计描述", 1)
    add_heading(doc, "2.1 项目背景", 2)
    for text in [
        "自学考试考籍管理覆盖考生基础信息、准考证号、考籍号、专业层次、报考批次、材料证明、免考申请、课程顶替、转入转出和毕业审核等多类数据。传统管理方式依赖线下材料、人工核验和分散表格，容易出现信息重复录入、材料状态不清、流程记录缺失、审核依据难追溯等问题。随着省级考试院业务量持续增长，考籍业务需要更加规范的信息系统来支撑材料留存、流程审批、权限控制和审计追踪。",
        "本项目面向上述痛点，建设省考试院自学考试考籍管理系统。系统不只实现简单的档案维护，还将免考、课程顶替、转考和毕业申请抽象为统一业务流程，通过通用申请表、审核记录表和扩展字段表保存差异化业务信息。同时，系统引入材料预处理、材料智能核验、智能问答、图像分类、目标检测、ASR 和 TTS 等辅助能力，为材料审核人员提供风险提示和复核建议。",
        "从课程实践角度看，本项目能够覆盖 Web 应用开发的完整链路：数据库建模、后端分层设计、前端工程化、接口联调、文件上传下载、权限拦截、日志审计、算法服务调用和测试文档编写。通过该项目，可以把软件工程中的需求分析、概要设计、详细设计、编码实现、测试验证和项目交付串联起来。"
    ]:
        add_para(doc, text)
    add_screenshot_placeholder(doc, "系统登录页面", "建议截图内容：登录页、账号密码输入框、登录按钮与系统名称。")

    add_heading(doc, "2.2 项目开发目的", 2)
    for text in [
        "功能层面，本项目旨在构建一个支持省考试院自学考试考籍管理的综合平台。系统需要帮助管理人员集中维护考生信息和考籍档案，实现材料上传、预览、下载、审核、申请提交、审核通过、驳回、撤回、流程记录查询、档案状态联动和日志审计等能力。通过统一的数据模型和接口规范，减少手工管理造成的数据不一致，提高业务办理透明度。",
        "技术层面，本项目旨在训练团队对前后端分离架构、Spring Boot 分层开发、Vue3 组件化开发、MyBatis-Plus 数据访问、MySQL 数据库设计、FastAPI 算法服务、Token 认证和接口安全的综合运用能力。团队需要在真实项目结构中完成接口开发、页面开发、联调测试和问题修复，积累端到端工程实践经验。",
        "实训层面，本项目强调协作开发与文档交付。团队在开发过程中遵守中文文档、关键代码 Doxygen 注释、禁止批量删除文件等项目约定，形成了接口说明、数据库设计、部署说明和测试用例等文档，便于教师验收和后续迭代维护。"
    ]:
        add_para(doc, text)

    add_heading(doc, "2.3 项目开发环境", 2)
    add_table(
        doc,
        ["类别", "工具或版本", "说明"],
        [
            ["操作系统", "Windows 10/11", "用于前后端开发、数据库管理和课程演示。"],
            ["后端环境", "JDK 17、Maven 3.8+、Spring Boot 3.3.5", "实现 REST 接口、业务服务、权限拦截和日志记录。"],
            ["前端环境", "Node.js 18+、Vue3、Vite、TypeScript、Element Plus", "实现管理端页面、路由、组件和接口调用。"],
            ["数据库", "MySQL 8.0", "保存系统用户、角色菜单、考生、考籍、材料、业务申请和审核记录。"],
            ["算法服务", "Python 3.10+、FastAPI", "提供材料预处理、智能问答、图像分析等辅助接口。"],
            ["开发工具", "IntelliJ IDEA、VS Code、Chrome、Git", "完成编码、调试、页面预览和版本管理。"],
        ],
    )
    add_code_block(
        doc,
        "后端启动：\ncd backend\nmvn spring-boot:run\n\n前端启动：\ncd frontend\nnpm install\nnpm run dev\n\n算法服务启动：\ncd algorithm-service\npython -m uvicorn app.main:app --host 127.0.0.1 --port 9000 --reload",
    )

    add_heading(doc, "2.4 系统总体设计", 2)
    for text in [
        "系统采用前后端分离和分层架构。前端负责页面展示、表单校验、路由控制、权限菜单和用户交互；后端负责统一认证、业务校验、流程处理、文件管理、日志审计和算法服务转发；数据库负责持久化考籍业务数据；算法服务通过内部 API Key 与后端通信，避免前端直接暴露算法服务密钥。",
        "后端工程按照 Controller、Service、Mapper、Entity、DTO、VO、Config、Common、Util 等包划分职责。Controller 负责接收请求和返回统一响应；Service 负责业务规则和事务边界；Mapper 负责数据库访问；DTO 用于接收前端请求；VO 用于返回页面展示数据；Config 存放认证、权限、跨域、日志和 MyBatis-Plus 配置。",
        "前端工程采用 Vue3 + TypeScript + Element Plus。路由配置集中在 router/index.ts 中，不同菜单绑定 permission 权限标识。登录成功后前端保存 Token 和用户权限，后续请求由 Axios 拦截器自动携带 Authorization 请求头，页面访问由路由守卫控制。"
    ]:
        add_para(doc, text)
    add_figure(doc, "diagram-architecture.png", "图：系统总体架构图")
    add_screenshot_placeholder(doc, "系统工作台页面", "建议截图内容：顶部欢迎面板、统计卡片、常用业务快捷入口和待办信息。")
    add_page_break(doc)

    add_heading(doc, "2.5 数据库设计", 2)
    for text in [
        "数据库采用 exam_record 作为业务库，表结构通过 sql/init.sql 初始化。系统表包括 sys_user、sys_role、sys_menu、sys_user_role、sys_role_menu、sys_login_log 和 sys_operation_log，用于认证授权与审计。考籍核心表包括 candidate、student_record、record_material、material_type、record_status_log 和 record_change_log，用于保存考生、档案、材料和状态变更。流程业务表包括 business_application、audit_record、process_status 和 application_extension_field，用于统一承载免考、顶替、转考和毕业申请。",
        "设计上，系统将不同流程业务的公共字段抽象到 business_application，例如申请编号、业务类型、考籍档案 ID、考生 ID、申请标题、当前状态、提交人、审核人和提交时间。不同业务的差异字段既保存到 extension_data_json 快照中，也可拆分写入 application_extension_field，便于列表检索和详情展示。审核动作统一写入 audit_record，用于业务详情页展示时间线。",
        "档案状态联动采用 record_status_log 和 record_change_log 双记录。前者关注档案状态从什么值变成什么值，后者关注字段级变更轨迹。转出审核通过后档案状态更新为 TRANSFERRED_OUT，毕业审核通过后档案状态更新为 GRADUATED。该设计能够满足后续审计和业务追溯需求。"
    ]:
        add_para(doc, text)
    add_figure(doc, "diagram-er.png", "图：核心数据关系图")
    add_table(
        doc,
        ["表类别", "主要表", "设计作用"],
        [
            ["系统权限", "sys_user、sys_role、sys_menu、sys_user_role、sys_role_menu", "支撑用户登录、角色授权、菜单权限和接口权限控制。"],
            ["日志审计", "sys_login_log、sys_operation_log", "记录登录成功失败和关键业务操作，便于追踪问题。"],
            ["考籍核心", "candidate、student_record", "保存考生基础信息和自考考籍档案。"],
            ["材料管理", "record_material、material_type", "保存材料文件、材料类型、审核状态和预览下载信息。"],
            ["流程业务", "business_application、audit_record、process_status、application_extension_field", "统一承载免考、课程顶替、转考、毕业等流程申请。"],
            ["状态追踪", "record_status_log、record_change_log", "记录档案状态与字段变更轨迹。"],
        ],
    )

    add_heading(doc, "2.6 接口与权限设计", 2)
    add_para(doc, "后端接口统一以 /api 为前缀，返回结构统一为 code、message、data。认证接口包括 /api/auth/login、/api/auth/logout 和 /api/auth/me；系统管理接口位于 /api/system/**；考籍业务接口位于 /api/candidates、/api/records、/api/materials、/api/exemptions、/api/course-replacements、/api/transfers、/api/graduations；智能辅助接口位于 /api/ai/**。")
    add_table(
        doc,
        ["模块", "接口范围", "权限控制要点"],
        [
            ["登录认证", "/api/auth/**", "登录接口放行，其他接口需要 Bearer Token。"],
            ["系统管理", "/api/system/users、/roles、/menus、/logs", "仅系统管理员可访问，菜单权限与接口权限一致。"],
            ["考籍档案", "/api/candidates、/api/records", "考籍管理员维护，审核员按权限查看。"],
            ["材料管理", "/api/materials、/api/material-types", "支持上传、预览、下载、审核和业务材料绑定。"],
            ["流程申请", "/api/exemptions、/api/course-replacements、/api/transfers、/api/graduations", "只有 SUBMITTED 状态允许修改、撤回和审核。"],
            ["智能辅助", "/api/ai/**", "后端统一携带内部密钥调用算法服务，前端不直连算法服务。"],
        ],
    )
    add_para(doc, "权限实现上，AuthInterceptor 从 Authorization 请求头提取 Bearer Token，调用 AuthService 完成校验，并将当前用户写入 AuthContextHolder。PermissionInterceptor 根据用户权限标识判断是否可访问接口；前端路由 meta.permission 与菜单权限保持一致，用户无权限访问时会回到工作台。")
    add_page_break(doc)


def add_module_implementation(doc: Document) -> None:
    add_heading(doc, "2.7 模块具体实现", 2)
    modules = [
        (
            "2.7.1 登录认证与权限控制模块",
            [
                "登录认证模块采用 Token 机制。用户提交账号和密码后，后端校验账号状态、密码哈希和角色权限，生成带有用户 ID、用户名、角色和权限集合的 Token 返回前端。前端将 Token 保存到本地，并通过 Axios 请求拦截器统一写入 Authorization 请求头。",
                "后端 AuthInterceptor 负责在请求进入业务控制器前完成 Token 校验。校验成功时写入当前用户上下文，校验失败时返回统一错误响应。系统还设计了 PermissionInterceptor，对需要权限的接口进行二次校验，避免仅靠前端菜单隐藏造成越权风险。",
                "前端路由守卫负责页面级访问控制。如果未登录访问内部页面，会跳转到登录页；如果用户没有某个菜单权限，会自动回到工作台。该模块保证系统在多角色场景下能够根据管理员、考籍管理员、审核员等角色展示不同功能入口。",
            ],
            "登录认证流程",
            "建议截图内容：登录成功后的工作台、侧边菜单和用户信息区域。",
        ),
        (
            "2.7.2 工作台与统计模块",
            [
                "工作台是系统登录后的第一屏，主要展示系统概览、待处理材料数量、考籍档案数量、业务申请数量和常用业务入口。后端 DashboardService 汇总考籍、材料和流程申请数据，返回 DashboardStatsVO，前端用统计卡片和快捷入口呈现。",
                "页面使用 Element Plus 栅格和卡片组件构建，结合权限判断动态展示智能核验、考籍档案、材料审核、免考管理、课程顶替、转入转出、毕业管理和系统管理等入口。工作台减少了业务人员在多个菜单之间来回查找的时间。",
                "统计模块还承担验收演示入口的作用。演示时可以先进入工作台查看总体数据，再按快捷入口进入具体业务页面，形成从概览到明细的演示路径。",
            ],
            "工作台统计卡片",
            "建议截图内容：统计卡片、快捷入口和待审核材料提示。",
        ),
        (
            "2.7.3 考生与考籍档案模块",
            [
                "考生管理模块负责维护考生姓名、身份证号、准考证号、联系方式等基础信息。新增或修改考生时，后端需要校验身份证号和准考证号唯一，避免同一考生重复建档。系统还提供导入预览能力，用于批量数据导入前发现格式错误和重复记录。",
                "考籍档案模块以 student_record 表为核心，保存考籍号、考生 ID、专业、层次、批次、档案状态、归档标记和备注等信息。档案状态包括 NORMAL、SUSPENDED、TRANSFERRED_OUT、GRADUATED、ARCHIVED 等业务含义。每次状态调整都会写入状态日志和变更记录。",
                "前端页面提供分页查询、详情查看、新增、编辑、状态更新和归档操作。后端 StudentRecordService 负责业务校验，例如考籍号唯一、档案存在性校验、归档不能重复执行、状态变更必须记录原因等。",
            ],
            "考生与考籍档案列表",
            "建议截图内容：考生列表、考籍档案列表或档案详情页。",
        ),
        (
            "2.7.4 材料管理与材料审核模块",
            [
                "材料管理模块支持 jpg、jpeg、png、pdf 等材料上传。上传接口接收 multipart/form-data，校验 recordId、materialType 和文件格式后，将文件保存到后端本地材料目录，并在 record_material 表中写入材料记录。材料可以按档案查询，也可以绑定到具体业务申请。",
                "材料审核页面支持按业务编号或考生考籍查询相关业务，并展示该业务下的材料清单。审核员可以预览材料、下载材料、上传补充材料并完成材料审核。审核状态从 PENDING 变为 APPROVED 或 REJECTED 后，会影响毕业资格校验和业务申请审核。",
                "该模块与智能辅助模块联动。前端可以先触发材料预处理，展示格式校验、清晰度、分类候选和建议动作，再触发申请材料智能核验，获取必交材料、缺失材料、异常材料和人工复核建议。",
            ],
            "材料审核页面",
            "建议截图内容：业务查询表单、材料列表、预览按钮和审核结果。",
        ),
        (
            "2.7.5 免考申请模块",
            [
                "免考申请模块面向考生已获得其他课程或证书成绩后申请免考的场景。业务数据写入 business_application，业务类型为 EXEMPTION。申请字段包括 recordId、courseCode、courseName、sourceCourseCode、sourceCourseName、exemptionReason、materialIds 和 remark。",
                "后端提交申请前会校验考籍档案是否存在，材料是否属于当前档案，同一考籍档案同一课程是否存在在办或已通过申请。申请状态为 SUBMITTED 时允许修改、撤回、审核通过或驳回；进入 APPROVED、REJECTED、WITHDRAWN 后成为终态。",
                "流程动作统一写入 audit_record。前端详情页可按时间线展示提交、修改、撤回、通过和驳回记录，业务办理过程清晰可追溯。",
            ],
            "免考申请详情",
            "建议截图内容：免考申请表单、审核按钮和流程记录。",
        ),
        (
            "2.7.6 课程顶替模块",
            [
                "课程顶替模块包括规则维护和申请办理两部分。规则维护用于配置可顶替课程、来源课程、适用专业、适用层次、有效期和启用状态。申请办理时，系统根据规则校验考生是否满足顶替条件。",
                "课程顶替申请同样复用 business_application 和 audit_record。不同于免考，课程顶替更强调规则有效期、专业层次匹配和来源课程对应关系。规则失效或不适用时，后端会返回业务校验错误，避免审核员处理无效申请。",
                "前端页面需要同时支持规则管理和申请列表，保证管理员可维护规则，业务人员可提交和审核申请。该模块体现了将业务规则从代码中抽离到数据库配置的思想。",
            ],
            "课程顶替规则与申请",
            "建议截图内容：规则列表、申请列表或审核弹窗。",
        ),
        (
            "2.7.7 考籍转入转出模块",
            [
                "考籍转入转出模块处理考生跨省转考场景。转入申请需要填写来源省份、原考籍信息和转入说明；转出申请需要填写目标省份、转出原因和相关证明材料。申请状态流转与免考流程保持一致。",
                "该模块的重要设计是档案状态联动。转出申请审核通过后，系统调用 RecordStatusLinkageService 将 student_record.record_status 更新为 TRANSFERRED_OUT，同时写入 record_status_log 和 record_change_log。这样后续业务在查询档案时可以识别该档案已经转出。",
                "流程记录查询接口支持按 businessType、businessId、applicationId 或 recordId 查询，既可查看单个转考申请的审核轨迹，也可查看某个档案全部业务轨迹。",
            ],
            "考籍转入转出审核",
            "建议截图内容：转出申请列表、审核通过后档案状态变化。",
        ),
        (
            "2.7.8 毕业申请模块",
            [
                "毕业申请模块提供毕业资格校验、申请提交、修改、撤回、审核、驳回、结果查询和流程记录查询。提交毕业申请前，系统会检查考籍状态、档案基础信息、考生身份信息、毕业申请材料和材料审核状态。",
                "毕业资格校验返回 eligible、passedItems、failedItems 和 warningItems。当前版本要求档案状态为 NORMAL，考籍号、专业名称、层次、考生姓名和身份证号不为空，存在 GRADUATION 类型材料，并且毕业申请材料均已审核通过。",
                "审核通过毕业申请后，系统将档案状态联动为 GRADUATED，并写入状态日志和变更记录。该流程体现了业务申请与档案生命周期的关联，避免毕业完成后档案仍被当作正常在籍处理。",
            ],
            "毕业资格校验",
            "建议截图内容：毕业申请页面、资格校验结果和审核结果。",
        ),
        (
            "2.7.9 智能辅助与算法服务模块",
            [
                "智能辅助模块由后端 AiAssistController、AiAssistServiceImpl 和 Python FastAPI 算法服务组成。前端所有智能请求都先发给 Spring Boot 后端，后端统一携带 X-Internal-Api-Key 调用算法服务，避免浏览器端暴露内部密钥。",
                "算法服务提供材料预处理、申请材料智能核验、图像分类、目标检测、图像分割、智能问答、ASR 和 TTS 等接口。当前实现中，材料预处理会返回格式校验、清晰度检测、材料分类候选、建议动作和人工复核标识；申请材料智能核验会根据业务类型返回必交材料、缺失材料和异常材料。",
                "后端对算法服务异常进行统一封装。当算法服务不可用、密钥不一致或参数错误时，后端返回统一业务错误，不向前端暴露内部堆栈。这样既提升了系统安全性，也让前端错误处理更加稳定。",
            ],
            "智能辅助页面",
            "建议截图内容：材料智能核验结果、智能问答或图像分析结果。",
        ),
        (
            "2.7.10 日志管理与审计模块",
            [
                "日志管理模块分为登录日志和操作日志。登录日志记录账号、登录状态、失败原因、客户端 IP 和登录时间；操作日志记录模块名称、操作类型、请求方式、请求地址、查询参数、响应摘要、操作状态、错误信息、操作人、客户端 IP、操作时间和耗时。",
                "系统通过 OperationLogInterceptor 和 OperationLogAdvice 自动记录已认证的 /api/** 内部接口操作，排除登录接口、健康检查接口、日志查询接口和静态资源，避免产生审计噪音。",
                "日志页面支持分页查询和详情查看。该模块为项目验收和故障排查提供依据，也符合真实管理系统对操作留痕和责任追溯的要求。",
            ],
            "日志管理页面",
            "建议截图内容：登录日志列表、操作日志列表或日志详情。",
        ),
    ]
    for heading, paragraphs, shot_title, shot_note in modules:
        add_heading(doc, heading, 3)
        for text in paragraphs:
            add_para(doc, text)
        add_screenshot_placeholder(doc, shot_title, shot_note)
        if heading in {"2.7.3 考生与考籍档案模块", "2.7.6 课程顶替模块", "2.7.9 智能辅助与算法服务模块"}:
            add_page_break(doc)


def add_detailed_design_and_integration(doc: Document) -> None:
    add_heading(doc, "2.8 关键设计与联调说明", 2)
    add_heading(doc, "2.8.1 统一业务申请模型设计", 3)
    for text in [
        "在考籍管理业务中，免考申请、课程顶替申请、考籍转入转出申请和毕业申请虽然表单字段不同，但它们在流程层面具有明显共性：都需要关联考生和考籍档案，都需要申请编号，都有提交、修改、撤回、审核通过、审核驳回等动作，都需要保存审核意见和时间线。因此项目没有为每一类业务完全独立设计流程表，而是采用 business_application 作为统一申请主表。",
        "统一业务申请模型的优势在于可以减少重复代码。分页查询、详情查询、状态校验、流程记录写入、申请材料读取等基础逻辑可以在不同业务服务中复用；差异化字段通过 extension_data_json 和 application_extension_field 保存，使系统既具备统一流程能力，又不会丢失各业务自身字段。该设计也便于后续新增其他自考业务，例如成绩复核、专业变更、照片变更等。",
        "在实现时，各业务 Service 仍保留独立入口，避免所有逻辑堆积在一个巨大服务类中。免考服务关注免考课程和证明来源课程；课程顶替服务关注规则匹配；转考服务关注来源省份或目标省份；毕业服务关注资格校验和毕业批次。公共部分保持一致，差异部分独立封装，使代码结构更清晰。",
    ]:
        add_para(doc, text)
    add_table(
        doc,
        ["业务类型", "公共字段", "差异字段", "终态处理"],
        [
            ["免考申请", "申请编号、考籍档案、考生、状态、材料、审核人", "免考课程、证明来源课程、免考原因", "通过后保留申请结果，不直接改变档案状态。"],
            ["课程顶替", "申请编号、考籍档案、考生、状态、材料、审核人", "目标课程、来源课程、适用规则、顶替说明", "通过后形成课程顶替记录，便于毕业资格判断。"],
            ["考籍转出", "申请编号、考籍档案、考生、状态、材料、审核人", "目标省份、转出原因、转出说明", "通过后联动档案状态为 TRANSFERRED_OUT。"],
            ["毕业申请", "申请编号、考籍档案、考生、状态、材料、审核人", "毕业批次、学位申请类型、资格校验结果", "通过后联动档案状态为 GRADUATED。"],
        ],
    )
    add_figure(doc, "diagram-business-flow.png", "图：业务办理流程图")

    add_heading(doc, "2.8.2 状态机与流程记录设计", 3)
    for text in [
        "流程状态是本项目中最容易产生错误的部分。为避免终态数据被重复修改，系统约定 SUBMITTED 为可处理状态，APPROVED、REJECTED、WITHDRAWN 为终态。只有 SUBMITTED 状态允许修改、撤回、审核通过或审核驳回。后端每个业务服务都会在执行动作前检查当前状态，如果状态不满足条件，则抛出业务异常并返回统一错误响应。",
        "流程记录通过 audit_record 表保存。每次提交、修改、撤回、通过和驳回都会写入一条记录，记录中包含业务类型、业务 ID、申请 ID、档案 ID、审核动作、变更前状态、变更后状态、审核意见、操作人和操作时间。这样前端不仅可以展示当前状态，还可以按时间顺序展示完整办理轨迹。",
        "状态流转设计让系统具有较强的可解释性。审核员在页面上看到的不只是一个结果，而是每一步是谁操作、何时操作、为什么操作。对于考试院这类需要严谨留痕的业务系统，流程记录比单纯的状态字段更重要。",
    ]:
        add_para(doc, text)
    add_table(
        doc,
        ["当前状态", "允许动作", "目标状态", "记录动作"],
        [
            ["新建表单", "提交", "SUBMITTED", "SUBMIT"],
            ["SUBMITTED", "修改", "SUBMITTED", "UPDATE"],
            ["SUBMITTED", "撤回", "WITHDRAWN", "WITHDRAW"],
            ["SUBMITTED", "审核通过", "APPROVED", "APPROVE"],
            ["SUBMITTED", "审核驳回", "REJECTED", "REJECT"],
            ["APPROVED/REJECTED/WITHDRAWN", "再次修改或审核", "不允许", "返回业务错误"],
        ],
    )
    add_figure(doc, "diagram-status-flow.png", "图：申请状态流转图")

    add_heading(doc, "2.8.3 前后端接口联调设计", 3)
    for text in [
        "接口联调阶段，团队先确认统一响应格式，再逐个模块对齐字段。后端统一返回 Result 结构，成功时 code 为 200，message 为操作成功，data 为具体业务数据；失败时返回业务错误码和明确提示。前端 API 层只处理统一结构，不需要在每个页面重复判断响应格式。",
        "前端接口文件按业务划分，例如 auth.ts、candidate.ts、record.ts、material.ts、exemption.ts、course.ts、transfer.ts、graduation.ts、ai.ts 和 system.ts。页面组件不直接拼接复杂请求，而是调用对应 API 方法。这样做能够减少页面代码中的重复逻辑，也便于后续修改接口路径或请求参数。",
        "联调中比较典型的问题是字段命名差异。前端通常使用驼峰字段，算法服务使用下划线字段，数据库使用下划线列名。项目通过 DTO、VO、MyBatis-Plus 驼峰映射和后端算法调用封装解决该问题，保证各层使用符合自身习惯的命名方式，同时保持数据语义一致。",
    ]:
        add_para(doc, text)
    add_code_block(
        doc,
        "统一响应结构示例：\n{\n  \"code\": 200,\n  \"message\": \"操作成功\",\n  \"data\": {\n    \"applicationStatus\": \"SUBMITTED\",\n    \"currentNodeName\": \"业务审核\"\n  }\n}\n\n认证请求头示例：\nAuthorization: Bearer <token>",
    )

    add_heading(doc, "2.8.4 文件上传与预览处理", 3)
    for text in [
        "材料上传是本系统区别于普通信息维护系统的重要能力。后端需要同时处理业务数据和文件数据：业务数据写入 record_material 表，文件数据保存到 uploads/materials 目录。上传时必须校验文件后缀、文件大小、考籍档案是否存在、材料类型是否启用，避免无效文件进入系统。",
        "下载接口使用 Content-Disposition: attachment，预览接口使用 Content-Disposition: inline。这样浏览器可以直接预览图片和 PDF，也可以在需要时下载原始材料。前端在材料审核页展示材料列表时，会根据材料 ID 调用预览地址，并提供下载按钮给审核人员使用。",
        "由于项目协作约定禁止批量删除文件，材料删除接口只处理单条材料记录及其对应的单个本地文件，不提供批量删除能力。测试删除类接口时也只使用临时新增数据，避免误删演示数据或真实材料。",
    ]:
        add_para(doc, text)

    add_heading(doc, "2.8.5 智能核验结果落地方式", 3)
    for text in [
        "智能核验不是替代人工审核，而是为人工审核提供参考。材料预处理接口会根据文件名、材料类型提示、业务场景和文件属性返回材料分类、清晰度、格式校验和建议动作。申请材料智能核验接口会根据业务类型判断必交材料是否齐全，并标记异常材料。",
        "前端展示智能核验结果时，应突出需要人工关注的问题，例如缺少身份证材料、成绩单材料置信度较低、上传类别与识别类别不一致、重复上传同类材料等。审核员可以根据这些提示决定是否通过材料审核。最终业务审核仍由人工点击通过或驳回完成。",
        "这种设计保证了智能辅助模块的可控性。算法服务可以提高效率，但不会绕过业务规则直接改变申请状态。后端服务仍是业务状态变更的唯一入口，符合考试管理类系统对安全和责任边界的要求。",
    ]:
        add_para(doc, text)
    add_page_break(doc)

    add_heading(doc, "2.9 部署运行与演示流程", 2)
    for text in [
        "系统部署分为数据库、算法服务、后端服务和前端应用四个部分。首先创建 MySQL 数据库并执行初始化脚本，保证表结构和演示数据存在；然后启动 FastAPI 算法服务，确认 /health 接口正常；再启动 Spring Boot 后端服务，确认 /api/health 正常；最后启动 Vue3 前端开发服务，通过浏览器访问 http://localhost:5173。",
        "推荐演示路线为：使用 admin 账号登录系统，进入工作台查看统计概览；进入系统管理查看用户、角色、菜单和日志；进入考生管理新增或查询考生；进入考籍档案查看档案详情和变更记录；进入材料审核上传并审核材料；依次演示免考、课程顶替、转入转出和毕业申请；最后进入智能辅助页面展示材料预处理和问答能力。",
        "在演示过程中，如果算法服务未启动，普通考籍业务仍可正常使用，只有智能辅助接口会返回服务不可用提示。该设计保证了核心业务和智能辅助之间的松耦合，避免算法服务异常影响整个系统演示。",
    ]:
        add_para(doc, text)
    add_table(
        doc,
        ["步骤", "操作", "验证点"],
        [
            ["1", "执行 sql/init.sql 和 sql/test-data.sql", "数据库 exam_record、基础表和演示账号存在。"],
            ["2", "启动 algorithm-service", "http://localhost:9000/health 返回 UP。"],
            ["3", "启动 backend", "http://localhost:8088/api/health 返回后端服务正常。"],
            ["4", "启动 frontend", "http://localhost:5173 可以打开登录页。"],
            ["5", "使用 admin/admin123 登录", "进入工作台并展示完整菜单。"],
            ["6", "执行材料审核和业务申请流程", "状态变化、流程记录和日志均可查询。"],
        ],
    )

    add_heading(doc, "2.10 安全性与异常处理设计", 2)
    for text in [
        "安全性方面，系统从登录认证、权限控制、参数校验、文件校验、日志审计和内部服务密钥几个层面进行防护。登录成功后使用 Token 识别用户身份，后端不信任前端传入的用户信息，而是通过 Token 解析当前用户。所有内部业务接口都需要认证，系统管理和业务审核接口还需要权限标识校验。",
        "参数校验方面，前端会进行非空、格式和长度校验，后端仍会进行二次校验，防止绕过前端直接请求接口。对于不存在的 ID、重复业务数据、终态重复审核、材料不属于当前档案等情况，后端都会返回明确的业务错误。",
        "算法服务调用方面，后端通过 X-Internal-Api-Key 调用 FastAPI 服务。算法服务只信任携带正确内部密钥的请求。当前端请求智能辅助能力时，只能访问 Spring Boot 后端封装接口，不能直接拿到算法服务密钥。",
        "日志审计方面，系统记录登录日志和操作日志。登录失败、接口异常和关键业务操作都可以在日志管理页面查询，为后续排查问题和审计追踪提供依据。",
    ]:
        add_para(doc, text)
    add_table(
        doc,
        ["风险点", "处理方式", "效果"],
        [
            ["未登录访问", "AuthInterceptor 校验 Token", "未认证请求返回 401。"],
            ["越权访问", "权限标识与角色菜单绑定", "无权限接口返回 403 或被前端路由拦截。"],
            ["重复提交", "状态校验与唯一性校验", "避免重复在办申请和重复审核。"],
            ["文件风险", "限制文件类型和大小", "减少非法文件上传。"],
            ["算法服务泄密", "后端统一携带内部 API Key", "前端不暴露算法服务密钥。"],
            ["问题追溯", "登录日志和操作日志", "关键操作可查询、可定位。"],
        ],
    )
    add_page_break(doc)


def add_tests_and_summary(doc: Document) -> None:
    add_heading(doc, "三、系统测试与运行验证", 1)
    add_para(doc, "系统测试围绕登录认证、权限控制、系统管理、考籍核心业务、业务流程、智能辅助、异常输入和三端联调展开。测试前先执行 sql/init.sql 创建表结构，再执行 sql/test-data.sql 导入演示账号和基础数据。随后依次启动 MySQL、算法服务、后端服务和前端服务，使用浏览器访问 http://localhost:5173 进行页面验证。")
    add_table(
        doc,
        ["测试类别", "测试内容", "预期结果"],
        [
            ["登录认证", "使用 admin、record_manager、auditor 登录，测试错误密码和未登录访问。", "正确账号登录成功，错误密码返回 401，未登录访问被拦截。"],
            ["权限控制", "不同角色访问系统管理、材料审核、业务审核和智能辅助页面。", "菜单和接口权限一致，无权限接口返回 403 或被前端路由拦截。"],
            ["考籍业务", "新增考生、创建档案、更新状态、归档、查看变更记录。", "唯一性校验生效，状态日志和变更记录完整。"],
            ["材料管理", "上传图片/PDF、预览、下载、审核、按业务绑定材料。", "材料保存成功，预览下载正常，审核状态正确变化。"],
            ["流程申请", "提交免考、顶替、转考、毕业申请，执行修改、撤回、审核和驳回。", "状态流转合法，终态不可重复操作，流程记录完整。"],
            ["智能辅助", "调用材料预处理、申请材料智能核验、智能问答等接口。", "算法服务可用时返回识别结果，不可用时返回统一错误。"],
            ["异常场景", "缺少必填字段、非法 ID、重复业务、终态重复审核、算法服务不可用。", "返回清晰业务错误，不产生脏数据。"],
        ],
    )
    add_screenshot_placeholder(doc, "系统测试用例执行结果", "建议截图内容：接口测试、页面联调或测试记录表。")
    add_para(doc, "在测试过程中，重点验证了统一响应格式、流程状态约束、材料归属校验和档案状态联动。对于转出和毕业这类会改变档案生命周期的业务，测试时不仅检查业务申请状态，还检查 student_record、record_status_log 和 record_change_log 是否同步写入。对于智能辅助接口，测试时同时覆盖算法服务正常、密钥缺失、参数错误和服务不可用场景，保证后端能够稳定兜底。")
    add_code_block(
        doc,
        "后端测试：\ncd backend\nmvn test\n\n算法服务测试：\ncd algorithm-service\npython -m compileall app\npython -m pytest\n\n前端生产构建：\ncd frontend\nnpm run build",
    )
    add_page_break(doc)

    add_heading(doc, "四、项目问题与解决方案", 1)
    problems = [
        ("前后端权限不一致", "问题表现为前端隐藏了菜单，但用户仍可能直接请求后端接口。解决方案是在后端增加权限拦截器，前端路由 meta.permission 与后端权限标识保持一致。"),
        ("业务流程字段差异较多", "免考、顶替、转考和毕业申请字段不完全相同，如果每类业务单独建主表会增加开发成本。解决方案是设计通用申请表和扩展字段表，公共字段统一保存，差异字段以 JSON 快照和扩展字段保存。"),
        ("材料与业务绑定复杂", "材料既属于考籍档案，又可能用于某个业务申请。解决方案是在上传和业务查询时校验 recordId 与 applicationId 关系，保证材料不会被错误绑定到其他考生档案。"),
        ("档案状态联动容易遗漏", "转出和毕业审核通过后，如果只修改申请状态而不修改档案状态，会造成业务数据不一致。解决方案是抽取 RecordStatusLinkageService，统一处理状态更新、状态日志和变更记录。"),
        ("算法服务不可用影响页面体验", "智能辅助能力依赖 FastAPI 服务和内部密钥，如果服务未启动会导致调用失败。解决方案是后端捕获 RestClientException 并返回统一业务错误，前端展示可理解的失败提示。"),
        ("中文文档和注释规范", "项目要求所有文档使用中文，关键代码使用 Doxygen 注释。解决方案是在核心服务、拦截器和控制器类上补充 @brief、@details、@param 等说明，并统一整理 docs 目录文档。"),
    ]
    add_table(doc, ["问题", "解决方案"], [[p, s] for p, s in problems])
    for p, s in problems:
        add_para(doc, f"{p}：{s}")
    add_page_break(doc)

    add_heading(doc, "五、实训收获与心得", 1)
    reflections = [
        "通过本次省考试院自学考试考籍管理系统实训，我完整经历了一个业务管理系统从需求分析到交付文档的全过程。与平时单独完成某个实验相比，本项目的复杂度更高，因为它需要同时考虑数据模型、权限边界、流程状态、材料文件、页面交互、算法服务和测试验证。开发过程中我明显感受到，只有前期把业务对象和状态流转梳理清楚，后期编码和联调才不会反复返工。",
        "在后端开发方面，我进一步理解了 Spring Boot 分层架构的价值。Controller 不应承载过多业务逻辑，Service 需要负责业务校验和事务一致性，Mapper 只处理数据访问，DTO 和 VO 则让请求参数和响应结构更加清晰。通过实现认证拦截、权限判断、统一异常处理和统一响应格式，我对企业级后台系统的基础能力有了更系统的认识。",
        "在前端开发方面，我认识到管理系统页面不仅要能展示数据，还要帮助用户高效完成工作。比如材料审核页面需要支持按业务编号和考生考籍两种方式查询，工作台需要提供待办提醒和快捷入口，流程详情需要让审核员快速看到当前状态和历史记录。组件化开发和 TypeScript 类型约束提高了页面维护性，也降低了联调时字段写错的概率。",
        "在数据库设计方面，我学会了将业务流程抽象成通用模型。免考、课程顶替、转考和毕业看起来是不同业务，但它们都有申请编号、申请人、档案、状态、审核人、审核时间和流程记录。将这些共性抽象出来后，系统的扩展性明显提升，后续新增其他考籍业务时可以复用同一套流程基础表。",
        "在智能辅助模块中，我体会到算法服务与业务系统之间需要清晰边界。算法服务负责识别和建议，最终业务决策仍应由后端规则和人工审核共同完成。后端统一封装算法调用、密钥和异常处理，既保护了内部服务，也让前端可以用稳定接口完成展示。",
        "本次实训也让我意识到文档和测试的重要性。接口说明、数据库设计、部署说明和测试用例不是附属工作，而是保证团队协作和项目验收的基础。尤其在多模块联调时，清晰的接口文档能够减少沟通成本；完整的测试用例能够帮助我们快速定位问题是否发生在前端、后端、数据库还是算法服务。",
    ]
    for text in reflections:
        add_para(doc, text)
    add_page_break(doc)

    add_heading(doc, "六、总结与展望", 1)
    for text in [
        "本项目完成了省考试院自学考试考籍管理系统的主要功能，实现了登录认证、权限管理、考生管理、考籍档案、材料管理、免考申请、课程顶替、考籍转入转出、毕业申请、日志审计和智能辅助等模块。系统采用 Spring Boot + MyBatis-Plus + MySQL + Vue3 + FastAPI 的技术路线，具备较完整的前后端分离架构和三端联调能力。",
        "从实现效果看，系统能够支撑课程演示中的主业务流程：管理员登录后进入工作台，维护考生和考籍档案，上传并审核材料，提交免考、顶替、转考或毕业申请，审核员完成审核，系统记录流程轨迹并根据业务规则联动档案状态。智能辅助模块可以在材料审核中提供格式、清晰度、分类和缺失材料建议，提高审核效率。",
        "后续若继续完善项目，可以从以下方向迭代：一是增加更完善的数据导入导出和批量校验能力；二是将材料文件存储迁移到对象存储或专用文件服务；三是完善流程引擎能力，使业务状态流转可配置；四是接入更真实的 OCR 和图像识别模型，提高材料核验准确率；五是增加自动化测试、接口契约测试和前端端到端测试，提高系统稳定性。",
    ]:
        add_para(doc, text)
    add_table(
        doc,
        ["后续优化方向", "具体内容"],
        [
            ["数据治理", "完善考生导入、重复数据合并、字段标准化和数据质量报告。"],
            ["流程配置", "将流程节点、审核角色和状态流转从代码中抽离为可配置规则。"],
            ["智能核验", "接入 OCR、证件识别和更准确的图像质量检测模型。"],
            ["安全审计", "增强密码策略、Token 黑名单、敏感字段脱敏和操作审计报表。"],
            ["自动测试", "补充后端单元测试、接口测试、前端构建检查和页面端到端测试。"],
        ],
    )

    add_page_break(doc)
    add_heading(doc, "附录一：核心接口清单", 1)
    add_table(
        doc,
        ["模块", "接口", "说明"],
        [
            ["认证", "POST /api/auth/login", "用户登录并返回 Token。"],
            ["认证", "GET /api/auth/me", "获取当前登录用户信息和权限。"],
            ["工作台", "GET /api/dashboard/stats", "获取首页统计数据。"],
            ["考生", "GET/POST/PUT/DELETE /api/candidates", "考生分页、新增、修改、删除和导入预览。"],
            ["考籍", "GET/POST/PUT /api/records", "考籍档案查询、创建、更新、状态变更和归档。"],
            ["材料", "POST /api/materials/upload", "上传档案材料。"],
            ["材料", "GET /api/materials/{id}/preview", "预览材料文件。"],
            ["免考", "POST /api/exemptions", "提交免考申请。"],
            ["顶替", "POST /api/course-replacements/applications", "提交课程顶替申请。"],
            ["转考", "POST /api/transfers", "提交转入或转出申请。"],
            ["毕业", "GET /api/graduations/eligibility/{recordId}", "毕业资格校验。"],
            ["智能", "POST /api/ai/material-preprocess", "材料预处理。"],
            ["智能", "POST /api/ai/applications/{applicationId}/material-audit", "按业务申请核验材料。"],
            ["日志", "GET /api/system/logs/operation/page", "查询操作日志。"],
        ],
    )
    add_heading(doc, "附录二：截图占位清单", 1)
    add_table(
        doc,
        ["序号", "截图名称", "建议替换位置"],
        [[str(i + 1), name, "正文对应截图占位符处"] for i, name in enumerate([
            "系统登录页面", "系统工作台页面", "登录认证流程", "工作台统计卡片", "考生与考籍档案列表",
            "材料审核页面", "免考申请详情", "课程顶替规则与申请", "考籍转入转出审核", "毕业资格校验",
            "智能辅助页面", "日志管理页面", "系统测试用例执行结果"
        ])],
    )


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("省考试院自学考试考籍管理系统实训报告")
        set_run_font(r, size=9)


def main() -> None:
    make_diagrams()
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    for style_name, size in [("Normal", 11), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12), ("Heading 4", 11)]:
        style = styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        if style_name != "Normal":
            style.font.bold = True

    build_cover(doc)
    add_responsibility(doc)
    add_design_description(doc)
    add_module_implementation(doc)
    add_detailed_design_and_integration(doc)
    add_tests_and_summary(doc)
    add_footer(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
