from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = next((ROOT / "docs").glob("10.*12*.docx"))
SQL = ROOT / "sql" / "init.sql"
BACKUP = ROOT / "target" / "report-before-db-details.docx"


@dataclass
class Column:
    name: str
    type_and_constraint: str
    comment: str


@dataclass
class TableInfo:
    name: str
    comment: str
    columns: list[Column]
    indexes: list[str]


def set_run_font(run, size: float = 10.5, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(p, first: bool = True, size: float = 10.5) -> None:
    p.paragraph_format.first_line_indent = Cm(0.74) if first else Pt(0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_run_font(run, size=size)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "9E9E9E")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, header: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    set_run_font(r, size=8.5 if not header else 9, bold=header)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        shade_cell(cell, "F2F4F7")


def parse_tables(sql_text: str) -> list[TableInfo]:
    pattern = re.compile(
        r"CREATE TABLE IF NOT EXISTS\s+(\w+)\s*\((.*?)\)\s*COMMENT='([^']*)';",
        re.S,
    )
    tables: list[TableInfo] = []
    for name, body, comment in pattern.findall(sql_text):
        columns: list[Column] = []
        indexes: list[str] = []
        for raw in body.splitlines():
            line = raw.strip().rstrip(",")
            if not line:
                continue
            upper = line.upper()
            if upper.startswith(("PRIMARY ", "UNIQUE ", "KEY ", "INDEX ", "CONSTRAINT ")):
                indexes.append(line)
                continue
            m = re.match(r"(\w+)\s+(.+?)(?:\s+COMMENT\s+'([^']*)')?$", line)
            if not m:
                continue
            columns.append(Column(m.group(1), m.group(2), m.group(3) or ""))
        tables.append(TableInfo(name, comment, columns, indexes))
    return tables


def find_marker(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(text):
            return p._p
    raise RuntimeError(f"未找到插入位置：{text}")


def add_paragraph_before(doc: Document, marker, text: str = "", style: str | None = None, first: bool = True, size: float = 10.5, bold: bool = False):
    p = doc.add_paragraph()
    marker.addprevious(p._p)
    if style:
        p.style = style
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    set_paragraph(p, first=first, size=size)
    return p


def add_heading_before(doc: Document, marker, text: str, level: int) -> None:
    p = add_paragraph_before(doc, marker, text, style=f"Heading {level}", first=False, size={2: 13, 3: 12, 4: 11}.get(level, 10.5), bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def add_page_break_before(doc: Document, marker) -> None:
    p = add_paragraph_before(doc, marker, first=False)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_table_before(doc: Document, marker, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    marker.addprevious(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, False)
    add_paragraph_before(doc, marker, "", first=False)


def add_placeholder_before(doc: Document, marker, title: str, note: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    marker.addprevious(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "FAFAFA")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(42)
    r = p.add_run(f"【数据库截图占位：{title}】\n{note}\n请在此处替换为实际数据库截图")
    set_run_font(r, size=12, bold=True, color="666666")
    add_paragraph_before(doc, marker, "", first=False)


def category_for(table: str) -> str:
    if table.startswith("sys_"):
        return "系统权限与日志"
    if table in {"candidate", "student_record", "record_material", "material_type", "record_status_log", "record_change_log"}:
        return "考籍核心数据"
    if table in {"business_application", "audit_record", "process_status", "application_extension_field"}:
        return "流程申请数据"
    if table == "course_replacement_rule":
        return "业务规则数据"
    return "其他业务数据"


def main() -> None:
    tables = parse_tables(SQL.read_text(encoding="utf-8"))
    doc = Document(str(DOCX))
    if any(p.text.strip().startswith("2.5.1 数据库表结构总览") for p in doc.paragraphs):
        print("已存在数据库字段明细，跳过重复插入。")
        return

    if not BACKUP.exists():
        BACKUP.parent.mkdir(parents=True, exist_ok=True)
        BACKUP.write_bytes(DOCX.read_bytes())

    marker = find_marker(doc, "2.6 接口与权限设计")

    add_heading_before(doc, marker, "2.5.1 数据库表结构总览", 3)
    add_paragraph_before(
        doc,
        marker,
        "本系统数据库名称为 exam_record，初始化脚本位于 sql/init.sql。数据库采用 utf8mb4 字符集，围绕系统权限、考籍核心数据、业务流程申请、材料管理、日志审计和智能辅助联调等场景设计。当前初始化脚本共创建 18 张业务表，既能支撑基础账号权限，又能支撑免考、课程顶替、考籍转入转出、毕业申请和材料审核等主流程。",
    )
    add_table_before(
        doc,
        marker,
        ["序号", "表名", "中文说明", "所属类别", "字段数"],
        [[str(i), t.name, t.comment, category_for(t.name), str(len(t.columns))] for i, t in enumerate(tables, 1)],
    )

    add_heading_before(doc, marker, "2.5.2 数据库截图占位", 3)
    add_paragraph_before(
        doc,
        marker,
        "下面预留数据库截图位置，后续可在 Navicat、DataGrip、MySQL Workbench 或命令行工具中截取真实数据库结构图、表列表和字段结构图后替换。建议截图时保持表名、字段名、字段类型和注释清晰可见。",
    )
    add_placeholder_before(doc, marker, "数据库表列表", "建议截图内容：exam_record 数据库下 18 张数据表列表。")
    add_placeholder_before(doc, marker, "数据库 ER 关系图", "建议截图内容：candidate、student_record、record_material、business_application、audit_record 等核心表关系。")
    add_placeholder_before(doc, marker, "考籍核心表字段", "建议截图内容：candidate、student_record、record_material 的字段结构。")
    add_placeholder_before(doc, marker, "流程业务表字段", "建议截图内容：business_application、audit_record、application_extension_field 的字段结构。")
    add_placeholder_before(doc, marker, "权限与日志表字段", "建议截图内容：sys_user、sys_role、sys_menu、sys_login_log、sys_operation_log 的字段结构。")

    add_heading_before(doc, marker, "2.5.3 数据表字段详细设计", 3)
    add_paragraph_before(
        doc,
        marker,
        "以下根据 sql/init.sql 中的 CREATE TABLE 语句整理各表字段。字段说明以数据库 COMMENT 为准，类型与约束列保留字段类型、是否非空、默认值、自增等信息，便于对照数据库截图和后端实体类。索引信息放在每张表字段表之后，体现唯一约束和常用查询条件。",
    )

    for index, table in enumerate(tables, 1):
        add_heading_before(doc, marker, f"2.5.3.{index} {table.name}（{table.comment}）", 4)
        add_paragraph_before(
            doc,
            marker,
            f"{table.name} 表用于{table.comment}相关数据存储，属于“{category_for(table.name)}”类别，共包含 {len(table.columns)} 个字段。该表与系统业务模块的关系应结合后端 Entity、Mapper 和 Service 使用场景理解。",
        )
        add_table_before(
            doc,
            marker,
            ["字段名", "类型与约束", "字段说明"],
            [[c.name, c.type_and_constraint, c.comment] for c in table.columns],
        )
        if table.indexes:
            add_table_before(
                doc,
                marker,
                ["索引或约束", "说明"],
                [[idx, "用于主键、唯一性约束或常用查询加速。"] for idx in table.indexes],
            )
        if index in {6, 12}:
            add_page_break_before(doc, marker)

    doc.save(str(DOCX))
    print(DOCX)


if __name__ == "__main__":
    main()
