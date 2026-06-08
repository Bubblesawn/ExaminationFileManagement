from __future__ import annotations

import textwrap
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = next((ROOT / "docs").glob("*实训报告.docx"))
ASSET_DIR = ROOT / "target" / "report-assets"
BACKUP = ROOT / "target" / "report-before-expand.docx"
ACTIVE_DOC: Document | None = None


CODE_SHOTS = [
    (
        "code-auth-interceptor.png",
        ROOT / "backend/src/main/java/com/exam/record/config/AuthInterceptor.java",
        17,
        88,
        "Token 认证拦截器核心代码",
    ),
    (
        "code-material-upload.png",
        ROOT / "backend/src/main/java/com/exam/record/service/impl/RecordMaterialServiceImpl.java",
        146,
        207,
        "档案材料上传与业务绑定核心代码",
    ),
    (
        "code-graduation-eligibility.png",
        ROOT / "backend/src/main/java/com/exam/record/service/impl/GraduationApplicationServiceImpl.java",
        187,
        266,
        "毕业资格校验与申请提交核心代码",
    ),
    (
        "code-ai-service.png",
        ROOT / "backend/src/main/java/com/exam/record/service/impl/AiAssistServiceImpl.java",
        202,
        246,
        "智能辅助服务调用核心代码",
    ),
    (
        "code-http-router.png",
        ROOT / "frontend/src/router/index.ts",
        1,
        55,
        "前端路由守卫与权限控制核心代码",
    ),
    (
        "code-material-view.png",
        ROOT / "frontend/src/views/material/MaterialAuditView.vue",
        340,
        567,
        "材料审核页面查询、上传和识别逻辑代码",
    ),
    (
        "code-fastapi-routes.png",
        ROOT / "algorithm-service/app/api/routes.py",
        23,
        83,
        "FastAPI 算法服务路由代码",
    ),
    (
        "code-image-analysis.png",
        ROOT / "algorithm-service/app/services/image_analysis_service.py",
        39,
        100,
        "材料图片质量分析代码",
    ),
]


def font(size: int, bold: bool = False, mono: bool = False):
    candidates = []
    if mono:
        candidates.extend([Path(r"C:\Windows\Fonts\consola.ttf"), Path(r"C:\Windows\Fonts\cour.ttf")])
    if bold:
        candidates.extend([Path(r"C:\Windows\Fonts\msyhbd.ttc"), Path(r"C:\Windows\Fonts\simhei.ttf")])
    candidates.extend([Path(r"C:\Windows\Fonts\msyh.ttc"), Path(r"C:\Windows\Fonts\simsun.ttc")])
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def make_code_screenshot(output_name: str, source: Path, start: int, end: int, title: str) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    lines = source.read_text(encoding="utf-8", errors="replace").splitlines()
    selected = lines[start - 1 : end]
    max_chars = 118
    rendered: list[tuple[int | None, str]] = []
    for number, line in enumerate(selected, start):
        chunks = textwrap.wrap(line.expandtabs(4), max_chars, replace_whitespace=False, drop_whitespace=False) or [""]
        for idx, chunk in enumerate(chunks[:3]):
            rendered.append((number if idx == 0 else None, chunk))
    rendered = rendered[:58]

    width = 1500
    line_height = 24
    top = 92
    height = top + len(rendered) * line_height + 42
    img = Image.new("RGB", (width, height), "#0F172A")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, width, 62), fill="#1F4E79")
    draw.text((28, 17), title, font=font(26, True), fill="#FFFFFF")
    draw.text((28, 68), source.as_posix().replace(ROOT.as_posix() + "/", ""), font=font(18), fill="#CBD5E1")
    mono = font(18, mono=True)
    y = top
    for number, code in rendered:
        if number is not None:
            draw.text((28, y), f"{number:>4}", font=mono, fill="#94A3B8")
        color = "#E2E8F0"
        stripped = code.strip()
        if stripped.startswith(("*", "/**", "*/", "//")):
            color = "#7DD3FC"
        elif stripped.startswith(("public ", "private ", "protected ", "def ", "function ", "async ")):
            color = "#FDE68A"
        elif stripped.startswith(("@", "import ", "from ")):
            color = "#C4B5FD"
        draw.text((95, y), code[:145], font=mono, fill=color)
        y += line_height
    img.save(ASSET_DIR / output_name)


def set_run_font(run, size: float = 11, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(p, size: float = 11, first: bool = True) -> None:
    p.paragraph_format.first_line_indent = Cm(0.74) if first else Pt(0)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_run_font(run, size=size)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "9E9E9E")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, header: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, 9.5 if not header else 10, header)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        shade_cell(cell, "F2F4F7")


def insert_after_marker(doc: Document, marker_text: str):
    marker = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(marker_text):
            marker = p
            break
    if marker is None:
        raise RuntimeError(f"未找到插入标记：{marker_text}")
    return marker._p


class Inserter:
    def __init__(self, marker_element):
        self.marker = marker_element

    def _insert(self, element):
        self.marker.addprevious(element)

    def paragraph(self, text: str = "", style: str | None = None, size: float = 11, bold: bool = False, first: bool = True):
        p = OxmlElement("w:p")
        self._insert(p)
        para = self.marker.getparent().xpath("./w:p")[-2]
        # xpath returns raw elements; wrap by using document object is awkward, so create via temporary paragraph below.
        return para


def insert_paragraph_before(marker_element, text: str, style: str | None = None, size: float = 11, bold: bool = False, first: bool = True):
    if ACTIVE_DOC is None:
        raise RuntimeError("未设置活动文档")
    para = ACTIVE_DOC.add_paragraph()
    marker_element.addprevious(para._p)
    if style:
        para.style = style
    para.paragraph_format.first_line_indent = Cm(0.74) if first else Pt(0)
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.space_after = Pt(4)
    if text:
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return para


def insert_heading(marker, text: str, level: int = 2) -> None:
    para = insert_paragraph_before(marker, text, style=f"Heading {level}", size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, first=False)
    para.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    para.paragraph_format.space_after = Pt(5)


def insert_page_break(marker) -> None:
    para = insert_paragraph_before(marker, "", first=False)
    para.add_run().add_break(WD_BREAK.PAGE)


def insert_image(marker, image_name: str, caption: str, width_cm: float = 14.2) -> None:
    image_path = ASSET_DIR / image_name
    if not image_path.exists():
        return
    para = insert_paragraph_before(marker, "", first=False)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image_path), width=Cm(width_cm))
    cap = insert_paragraph_before(marker, caption, size=9.5, bold=True, first=False)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def insert_table(marker, headers: list[str], rows: list[list[str]]) -> None:
    tbl = OxmlElement("w:tbl")
    marker.addprevious(tbl)
    from docx.table import Table

    table = Table(tbl, marker.getparent())
    table._tbl.get_or_add_tblPr()
    # Build table through python-docx after inserting a minimal tbl.
    table = Document().add_table(rows=1, cols=1)


def add_table_before(doc: Document, marker, headers: list[str], rows: list[list[str]]) -> None:
    temp = Document()
    table = temp.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell(row.cells[i], value, False)
    marker.addprevious(table._tbl)
    insert_paragraph_before(marker, "", first=False)


def build_expansion(doc: Document, marker) -> None:
    global ACTIVE_DOC
    ACTIVE_DOC = doc
    insert_page_break(marker)
    insert_heading(marker, "2.11 核心代码实现说明", 2)
    intro = [
        "为使报告内容更加贴合真实项目实现，本节补充系统关键代码的实现说明。项目代码按照后端、前端、算法服务三条主线组织：后端侧重点是认证拦截、业务校验、材料文件处理、流程状态写入和算法服务封装；前端侧重点是路由权限、页面状态管理、接口请求封装和材料审核交互；算法服务侧重点是内部鉴权、统一响应、材料预处理和图片质量分析。通过补充代码说明，可以更直观地体现团队在实训中完成的不只是页面展示，还包括完整的工程实现。",
        "本项目按照协作约定，在关键 Java 类和 Python 路由中补充了 Doxygen 风格注释。注释中使用 @brief 描述函数职责，使用 @details 说明业务规则，使用 @param 和 @return 描述输入输出。这样做能够让读者在阅读代码截图时快速理解模块边界，也方便后续维护者定位关键业务逻辑。",
        "以下代码截图均来自当前项目仓库，覆盖认证、材料、毕业、智能辅助、路由守卫、材料审核页面和算法服务等重点实现。报告中对每段代码的业务作用、实现步骤和异常处理方式进行了说明。",
    ]
    for text in intro:
        insert_paragraph_before(marker, text)

    add_table_before(
        doc,
        marker,
        ["代码位置", "所属层次", "主要作用"],
        [
            ["AuthInterceptor.java", "后端配置层", "统一校验 Bearer Token，写入用户上下文，认证失败时返回统一错误响应。"],
            ["RecordMaterialServiceImpl.java", "后端业务层", "处理材料上传、文件保存、材料入库、业务材料绑定和审核通过。"],
            ["GraduationApplicationServiceImpl.java", "后端业务层", "实现毕业资格校验、毕业申请提交、流程记录和档案状态联动。"],
            ["AiAssistServiceImpl.java", "后端集成层", "封装后端到 FastAPI 算法服务的调用，统一携带内部 API Key。"],
            ["router/index.ts", "前端路由层", "根据 Token 和权限标识控制页面访问。"],
            ["MaterialAuditView.vue", "前端页面层", "实现业务材料查询、上传、刷新、预览、下载和即时识别。"],
            ["routes.py", "算法接口层", "定义材料预处理、材料核验、图像识别和智能问答路由。"],
            ["image_analysis_service.py", "算法服务层", "读取图片并计算亮度、对比度、边缘得分和主体区域。"],
        ],
    )

    insert_heading(marker, "2.11.1 Token 认证拦截器实现", 3)
    for text in [
        "后端认证拦截器是系统安全边界的第一道关口。它在请求进入控制器之前执行，负责从 Authorization 请求头中提取 Bearer Token，并交由 AuthService 完成校验。如果校验通过，系统将当前用户信息写入 AuthContextHolder，后续 Service 层即可获取当前操作人，用于记录上传人、审核人和日志操作人。",
        "拦截器中特别放行 OPTIONS 请求，这是前后端分离项目常见的跨域预检请求。如果不放行，浏览器在真正发送业务请求前就会被拦截，导致前端页面出现跨域或认证异常。认证失败时，拦截器不直接抛出未处理异常，而是使用 Result.fail 生成统一 JSON 响应，使前端能按统一格式展示错误信息。",
        "afterCompletion 方法中清理 AuthContextHolder 是一个容易被忽略但非常重要的细节。由于 Web 容器会复用线程，如果请求结束后不清理 ThreadLocal 中的用户信息，后续请求可能读到上一次请求的用户上下文，造成安全隐患。",
    ]:
        insert_paragraph_before(marker, text)
    insert_image(marker, "code-auth-interceptor.png", "图：Token 认证拦截器核心代码")

    insert_heading(marker, "2.11.2 材料上传与业务绑定实现", 3)
    for text in [
        "材料上传模块同时处理文件流和业务数据。后端首先校验考籍档案是否存在，再校验材料类型是否启用，随后检查文件是否为空、后缀是否在白名单范围内。项目当前允许 jpg、jpeg、png 和 pdf，能够覆盖身份证、准考证、成绩单、照片和证明材料等常见材料。",
        "文件保存路径按 recordId、上传日期和 UUID 文件名组织。使用 UUID 可以避免原始文件名重复导致覆盖，同时保留 originalFileName 字段用于页面展示。代码中使用 normalize 后再判断 targetPath 是否以 uploadRootPath 开头，防止恶意文件名构造目录穿越路径。",
        "按业务编号上传材料时，系统先查询 business_application，取出申请关联的 recordId，再复用普通材料上传逻辑保存文件。上传成功后调用 appendBusinessMaterialId 将新材料 ID 写入申请材料列表。这样材料既属于考籍档案，也能同步绑定到具体业务申请，满足材料审核页面按业务查看的需要。",
    ]:
        insert_paragraph_before(marker, text)
    insert_image(marker, "code-material-upload.png", "图：档案材料上传与业务绑定核心代码")

    insert_heading(marker, "2.11.3 毕业资格校验与申请提交实现", 3)
    for text in [
        "毕业申请模块比普通流程申请更复杂，因为它不仅需要保存申请状态，还需要判断考生是否具备毕业资格。checkGraduationEligibility 方法不会修改数据库，只根据当前考籍档案、考生基础信息和毕业材料审核状态生成校验结果。因此它既可以被前端预校验接口调用，也可以在提交和修改毕业申请时复用。",
        "资格校验结果拆分为 passedItems、failedItems 和 warningItems。passedItems 记录已满足条件，例如考籍状态正常、身份信息完整、毕业材料已审核通过；failedItems 记录阻断提交的问题，例如缺少毕业材料或材料未审核通过；warningItems 用于提示但不一定阻断，例如档案已归档需要人工确认。",
        "提交毕业申请时，系统先执行同档案在办毕业申请唯一性校验，再校验材料归属，随后重新执行资格校验。只有资格校验通过后，才会创建 business_application，写入扩展字段快照，并写入 SUBMIT 流程记录。该流程保证毕业申请的业务数据、材料数据和流程记录能够保持一致。",
    ]:
        insert_paragraph_before(marker, text)
    insert_image(marker, "code-graduation-eligibility.png", "图：毕业资格校验与申请提交核心代码")

    insert_heading(marker, "2.11.4 智能辅助服务封装实现", 3)
    for text in [
        "智能辅助模块没有让前端直接访问算法服务，而是由 Spring Boot 后端统一封装调用。这样做有三个好处：第一，前端不暴露算法服务内部地址和 API Key；第二，后端可以把前端驼峰字段转换为算法服务需要的数据结构；第三，当算法服务不可用时，后端可以统一捕获异常并返回业务错误。",
        "AiAssistServiceImpl 中的 material-preprocess、application-material-audit、object-detect、image-segment、chat、asr 和 tts 等方法都复用 callAlgorithm。每次调用都携带业务 ID 和业务场景，便于后续扩展日志记录和链路追踪。",
        "按申请 ID 自动核验材料时，后端会读取 business_application 中的材料 ID 列表，再反查 record_material 获取文件地址、原始文件名和登记材料类型，最终组装 ApplicationMaterialAuditDTO 调用算法服务。这个封装使前端只需要传入 applicationId 即可完成一键核验，降低页面复杂度。",
    ]:
        insert_paragraph_before(marker, text)
    insert_image(marker, "code-ai-service.png", "图：智能辅助服务调用核心代码")
    insert_page_break(marker)

    insert_heading(marker, "2.12 前端关键实现说明", 2)
    insert_heading(marker, "2.12.1 路由守卫与权限菜单实现", 3)
    for text in [
        "前端路由配置将页面路径、组件和权限标识集中维护。每个内部页面都在 meta.permission 中声明需要的权限，例如 dashboard:view、record:view、material:audit:view、exemption:view、system:log:view 等。这样路由层、菜单层和后端权限层可以使用同一套权限标识，降低权限不一致风险。",
        "router.beforeEach 是前端页面访问控制的核心。用户访问登录页时直接放行；访问业务页面时，如果本地没有 Token，则跳转到登录页并携带 redirect 参数；如果用户有 Token 但没有目标页面权限，则回到工作台。该设计能避免用户通过地址栏直接访问未授权页面。",
        "前端权限控制只负责用户体验层面的引导，真正的安全控制仍由后端拦截器完成。即使用户绕过前端直接请求接口，后端仍会进行 Token 和权限校验。这种前后端双层控制是管理系统中较常见的设计方式。",
    ]:
        insert_paragraph_before(marker, text)
    insert_image(marker, "code-http-router.png", "图：前端路由守卫与权限控制核心代码")

    insert_heading(marker, "2.12.2 材料审核页面状态管理实现", 3)
    for text in [
        "材料审核页面是本项目交互较复杂的页面之一。它同时支持按业务编号查询和按考籍档案查询两种方式。按业务编号查询时，页面直接加载一个业务材料包；按考籍档案查询时，页面先列出该考生相关业务，再由审核员选择具体业务进行材料审核。",
        "页面中的 businessBundle 保存当前选中的业务申请材料包，recordBundles 保存按考籍档案查询出的业务列表，latestRecognition 保存最近一次材料识别结果。通过这些响应式状态，页面可以在查询、上传、刷新、预览、下载和审核动作之间保持数据同步。",
        "上传材料时，前端先执行文件格式和大小校验，再调用 uploadBusinessMaterial 上传材料。上传成功后，页面会同步刷新业务材料包，并对刚上传的图片材料执行即时分类识别。这样审核人员可以在上传后立刻看到算法给出的材料类别和置信度提示。",
    ]:
        insert_paragraph_before(marker, text)
    insert_image(marker, "code-material-view.png", "图：材料审核页面查询、上传和识别逻辑代码")

    insert_heading(marker, "2.12.3 前端异常提示与交互反馈", 3)
    for text in [
        "前端 API 封装统一处理后端 Result 响应。当后端返回 code 不等于 200 时，http.ts 会弹出错误提示并 reject 当前 Promise；当网络异常、登录凭证失效或服务不可用时，也会显示统一提示。页面组件只需要关注业务成功后的数据渲染，不需要在每个接口调用处重复解析响应格式。",
        "在材料审核页面中，查询、上传和远程搜索都设置了 loading 状态，避免用户连续点击造成重复请求。上传前校验文件后缀和大小，能够在浏览器端提前拦截不合规文件，减少后端压力。上传失败时通过 options.onError 反馈给 Element Plus 上传组件，页面可以恢复可操作状态。",
        "页面使用 Element Plus 的表格、表单、标签、按钮、上传、描述列表和空状态组件组合业务视图。不同业务状态使用不同 tag 类型展示，使审核员能够快速识别当前申请是已提交、审核通过、驳回还是撤回。",
    ]:
        insert_paragraph_before(marker, text)

    add_table_before(
        doc,
        marker,
        ["前端状态", "来源", "页面作用"],
        [
            ["businessBundle", "业务编号查询或考籍业务选择", "保存当前正在审核的业务申请和材料列表。"],
            ["recordBundles", "按考籍档案查询接口", "展示某个考生关联的免考、顶替、转考、毕业业务。"],
            ["latestRecognition", "材料即时识别接口", "展示最新上传材料的识别类别、置信度和建议。"],
            ["loading/uploading", "查询和上传动作", "控制按钮加载状态，避免重复提交。"],
            ["materialTypes", "材料类型维护接口", "提供上传材料类型下拉选择。"],
        ],
    )
    insert_page_break(marker)

    insert_heading(marker, "2.13 算法服务关键实现说明", 2)
    insert_heading(marker, "2.13.1 FastAPI 路由与内部鉴权", 3)
    for text in [
        "算法服务采用 FastAPI 实现，所有业务路由挂载在 /api 前缀下，并统一依赖 verify_internal_api_key。后端调用算法服务时会在请求头中携带 X-Internal-Api-Key，算法服务只处理内部密钥正确的请求。该设计避免外部浏览器或未授权客户端直接调用算法接口。",
        "routes.py 中定义了 image-classify、material-preprocess、application-material-audit、object-detect、image-segment、chat、asr 和 tts 等接口。接口层只负责参数接收和响应模型约束，具体业务处理委托给 mock_algorithm_service 或 image_analysis_service。这样算法服务结构清晰，后续替换真实模型时只需要改服务层实现。",
        "算法响应统一使用 AlgorithmResponse 模型，包含 code、message 和 data，与后端统一响应风格保持一致。后端收到算法服务响应后，可以直接包装给前端展示，也可以根据 code 和 data 做进一步业务判断。",
    ]:
        insert_paragraph_before(marker, text)
    insert_image(marker, "code-fastapi-routes.png", "图：FastAPI 算法服务路由代码")

    insert_heading(marker, "2.13.2 材料图片质量分析实现", 3)
    for text in [
        "材料预处理中的图片质量分析使用 PIL 读取本地图片，并计算亮度、对比度、边缘得分和主体区域。亮度用于判断低光和过曝风险，对比度用于判断材料是否发灰或文字不清，边缘得分用于粗略判断模糊风险，主体区域用于判断证件或纸张是否位于图片中。",
        "analyze_image 方法首先检查文件后缀是否属于支持图片类型，然后调用 resolve_image_path 将上传访问地址解析到算法服务可读取的本地路径。如果图片不存在或无法打开，方法不会抛出未处理异常，而是返回 loaded=false 和问题编码，方便上层服务生成可读提示。",
        "该实现虽然不是深度学习模型，但已经可以覆盖演示中的基础质量检查需求。后续如果接入 OCR 或目标检测模型，可以保留当前响应结构，在 data 中增加文字识别结果、证件区域坐标和真实性判断等字段。",
    ]:
        insert_paragraph_before(marker, text)
    insert_image(marker, "code-image-analysis.png", "图：材料图片质量分析代码")

    add_table_before(
        doc,
        marker,
        ["检测指标", "计算方式", "业务含义"],
        [
            ["brightness", "灰度均值 / 255", "判断图片是否过暗或过曝。"],
            ["contrast", "灰度标准差 / 128", "判断文字和背景是否区分明显。"],
            ["edge_score", "边缘滤波后均值 / 255", "粗略判断材料是否模糊。"],
            ["document_bbox", "基于灰度区域检测主体边界", "判断材料主体是否位于画面中。"],
            ["issues", "规则阈值输出问题编码", "供前端展示低清晰度、低对比度等提示。"],
        ],
    )
    insert_page_break(marker)

    insert_heading(marker, "2.14 关键业务流程详细说明", 2)
    flows = [
        (
            "2.14.1 材料上传与审核流程",
            [
                "材料上传流程从前端选择业务开始。审核员可以输入业务编号，也可以先按考籍档案搜索考生相关业务。页面拿到业务材料包后，展示业务基础信息、当前状态、已绑定材料和上传入口。",
                "用户选择材料类型并上传文件后，前端先校验文件后缀、MIME 类型和大小。校验通过后，后端保存文件并写入 record_material。随后，系统把材料 ID 追加到 business_application.material_ids_json，使该材料正式绑定到业务申请。",
                "审核员可以预览、下载、审核通过或删除单条材料。审核通过会写入审核状态、审核意见、审核人和审核时间。对于毕业申请而言，材料审核状态会直接影响毕业资格校验结果。",
            ],
        ),
        (
            "2.14.2 免考申请办理流程",
            [
                "免考申请要求考生提供目标免考课程、证明来源课程、免考原因和相关材料。提交前，后端会校验考籍档案是否存在，材料是否属于当前档案，以及同一档案同一课程是否已存在在办或已通过申请。",
                "免考申请提交后进入 SUBMITTED 状态。审核员可以根据材料和申请理由选择通过或驳回。通过后申请进入 APPROVED 终态，驳回后进入 REJECTED 终态。终态申请不允许再次修改或审核，避免流程结果被重复覆盖。",
                "每次提交、修改、撤回、通过和驳回都会写入 audit_record。前端在详情页展示流程时间线，使审核过程能够被追溯。",
            ],
        ),
        (
            "2.14.3 考籍转出状态联动流程",
            [
                "考籍转出申请审核通过后，不仅要更新申请状态，还要联动考籍档案状态。系统会将 student_record.record_status 更新为 TRANSFERRED_OUT，表示该考生档案已转出，后续业务不应继续按正常在籍处理。",
                "状态联动过程中，系统会同步写入 record_status_log 和 record_change_log。前者记录状态从 NORMAL 变为 TRANSFERRED_OUT 的过程，后者记录 recordStatus 字段变更轨迹。双记录设计能够同时满足业务状态查询和字段级审计。",
                "如果档案已经处于目标状态，系统不会重复写入状态日志，避免重复审核造成冗余记录。",
            ],
        ),
        (
            "2.14.4 毕业申请审核流程",
            [
                "毕业申请提交前会执行资格校验。校验项包括考籍状态是否正常、考生姓名和身份证号是否完整、是否存在毕业材料、毕业材料是否全部审核通过。校验失败时，系统返回失败项并阻止提交。",
                "毕业申请审核通过前，后端再次执行资格校验，避免材料或档案状态在提交后发生变化。审核通过后，申请状态变为 APPROVED，档案状态联动为 GRADUATED，同时写入流程记录、状态日志和变更记录。",
                "毕业申请结果查询只对已通过申请开放。这样可以避免用户在未完成审核时误认为已经取得毕业结果。",
            ],
        ),
    ]
    for heading, paragraphs in flows:
        insert_heading(marker, heading, 3)
        for text in paragraphs:
            insert_paragraph_before(marker, text)

    add_table_before(
        doc,
        marker,
        ["流程", "关键校验", "落库结果"],
        [
            ["材料上传", "档案存在、材料类型启用、文件格式合法、路径安全", "record_material 新增记录，业务申请材料 ID 列表更新。"],
            ["免考申请", "档案存在、材料归属、同课程无重复在办申请", "business_application 和 audit_record 写入。"],
            ["转出审核", "申请状态为 SUBMITTED、审核意见合法", "申请 APPROVED，档案状态 TRANSFERRED_OUT。"],
            ["毕业审核", "资格校验通过、材料审核通过", "申请 APPROVED，档案状态 GRADUATED。"],
            ["智能核验", "算法服务可用、内部密钥正确、材料可读取", "返回缺失材料、异常材料和建议动作。"],
        ],
    )
    insert_page_break(marker)

    insert_heading(marker, "2.15 联调问题记录与处理过程", 2)
    issue_rows = [
        ["登录后页面接口返回 401", "本地 Token 过期或后端重启导致令牌失效", "前端提示登录凭证失效，重新登录后恢复。"],
        ["材料上传后业务页面未刷新", "上传接口只返回材料记录，未同步业务材料包", "新增按业务编号上传接口，上传后返回最新 BusinessMaterialBundleVO。"],
        ["PDF 和图片预览方式不同", "浏览器对下载和预览响应头处理不同", "下载接口使用 attachment，预览接口使用 inline。"],
        ["毕业申请资格校验结果不清楚", "只返回布尔值无法说明失败原因", "拆分 passedItems、failedItems 和 warningItems。"],
        ["算法服务不可用时页面报错不友好", "后端未统一包装算法异常", "后端捕获 RestClientException 并返回统一业务错误。"],
        ["按考生查询材料无法定位具体业务", "一个考籍档案可能存在多个申请", "先展示业务列表，再由审核员选择业务材料包。"],
        ["文件名重复可能覆盖材料", "原始文件名直接保存存在冲突", "使用 UUID 作为存储文件名，保留原始文件名用于展示。"],
        ["路径穿越风险", "文件名可能包含 ../ 等特殊路径", "cleanPath、normalize，并校验 targetPath 位于 uploadRootPath 内。"],
    ]
    add_table_before(doc, marker, ["问题现象", "原因分析", "处理方式"], issue_rows)
    for row in issue_rows:
        insert_paragraph_before(marker, f"{row[0]}：{row[1]}。处理上，{row[2]} 该问题的处理过程使系统在联调阶段更加稳定，也让前端页面能够给出更清晰的业务反馈。")

    insert_heading(marker, "2.16 扩展测试用例补充", 2)
    add_table_before(
        doc,
        marker,
        ["用例编号", "测试目标", "操作步骤", "预期结果"],
        [
            ["TC-CODE-001", "验证 Token 拦截", "不携带 Authorization 请求 /api/records/page", "返回 401，业务控制器不执行。"],
            ["TC-CODE-002", "验证权限守卫", "普通审核员访问系统用户管理路由", "前端跳转工作台，后端接口拒绝越权访问。"],
            ["TC-CODE-003", "验证材料上传格式", "上传 exe 或 zip 文件作为材料", "前端或后端提示格式不支持。"],
            ["TC-CODE-004", "验证业务材料绑定", "按业务编号上传一份成绩单材料", "业务详情中 materialIds 包含新材料 ID。"],
            ["TC-CODE-005", "验证毕业资格阻断", "毕业材料未审核通过时提交毕业申请", "返回资格校验未通过，申请不入库。"],
            ["TC-CODE-006", "验证转出状态联动", "审核通过转出申请", "档案状态变为 TRANSFERRED_OUT，状态日志写入。"],
            ["TC-CODE-007", "验证智能核验缺失项", "免考申请只上传身份证材料", "返回成绩单和免考证明缺失提示。"],
            ["TC-CODE-008", "验证算法服务异常", "停止算法服务后调用 /api/ai/material-preprocess", "后端返回统一错误，不暴露内部堆栈。"],
            ["TC-CODE-009", "验证日志审计", "执行材料审核通过动作", "操作日志记录模块、操作人、请求地址和耗时。"],
            ["TC-CODE-010", "验证终态不可修改", "对 APPROVED 申请再次撤回或驳回", "返回状态不允许操作。"],
        ],
    )
    for text in [
        "扩展测试用例主要覆盖代码层面的关键分支，而不仅仅是页面是否能打开。认证、权限、文件、流程、算法和日志都属于系统风险较高的区域，需要在验收前重点测试。",
        "测试时应尽量使用 sql/test-data.sql 中的演示账号和临时新增数据。涉及删除和撤回的操作，不应直接作用于老师演示所需的基础数据，避免影响后续答辩展示。",
        "对于智能辅助测试，应分别覆盖算法服务正常和异常两种状态。正常状态下关注返回结构和识别建议，异常状态下关注后端错误封装和前端提示是否友好。",
    ]:
        insert_paragraph_before(marker, text)

    insert_page_break(marker)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for args in CODE_SHOTS:
        make_code_screenshot(*args)

    if not BACKUP.exists():
        BACKUP.write_bytes(DOCX.read_bytes())

    doc = Document(str(DOCX))
    marker = insert_after_marker(doc, "三、系统测试与运行验证")
    build_expansion(doc, marker)
    doc.save(str(DOCX))
    print(DOCX)


if __name__ == "__main__":
    main()
